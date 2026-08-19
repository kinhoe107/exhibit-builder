import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
OUT = Path(os.environ.get("EXHIBIT_GUIDED_SAMPLE_OUT", ROOT / "public" / "guided-sample"))
OUT.mkdir(parents=True, exist_ok=True)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5C6670"
LIGHT_BLUE = "E8EEF5"
CALLOUT = "F4F6F9"
CAUTION = "7A5A00"


def set_run_font(run, size=11, color="000000", bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)


def add_decimal_numbering(document):
    numbering = document.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    level.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    level.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "22")
    r_pr.append(size)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_number(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)


def create_statement():
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.add_run("EXHIBIT BUILDER | GUIDED SAMPLE"), 9, MUTED, True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    set_run_font(footer.add_run("Demonstration only | Guided sample"), 9, MUTED)

    kicker = document.add_paragraph()
    kicker.paragraph_format.space_after = Pt(4)
    set_run_font(kicker.add_run("GUIDED TUTORIAL"), 10, BLUE, True)
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    set_run_font(title.add_run("Guided Sample Witness Statement"), 25, INK, True)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    set_run_font(subtitle.add_run("This is an instructional guide, not a realistic witness statement."), 13, MUTED, False, True)

    metric = document.add_table(rows=1, cols=3)
    set_table_geometry(metric, [3120, 3120, 3120])
    for cell, label, value in zip(metric.rows[0].cells, ("STATEMENT", "EVIDENCE TYPES", "TEMPLATES"), ("1 DOCX guide", "PDF / DOCX / EML / XLSX", "Cover + full-table index")):
        shade(cell, LIGHT_BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(label + "\n"), 8.5, DARK_BLUE, True)
        set_run_font(p.add_run(value), 10, INK, True)

    callout = document.add_table(rows=1, cols=1)
    set_table_geometry(callout, [9360])
    shade(callout.cell(0, 0), CALLOUT)
    p = callout.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    set_run_font(p.add_run("How to use this sample: "), 10.5, DARK_BLUE, True)
    set_run_font(p.add_run("Run it from the front page, review the obvious matches, open the workbook-sheet step, inspect and confirm the supplied cover and index templates, group a few exhibits under an index heading, and use the unused checklist to practise adding an uncited exhibit."), 10.5, INK)

    heading = document.add_paragraph("What the tool should recognise", style="Heading 1")
    heading.paragraph_format.keep_with_next = True
    num_id = add_decimal_numbering(document)
    paragraphs = [
        "The witness statement supplied to Exhibit Builder must be a DOCX file. Evidence exhibits can be PDF, DOCX, EML or XLSX files. This first example is a PDF: I refer to the SAMPLE AGREEMENT dated 1 August 2026 [Exhibit].",
        "The tool recognises general placeholders labelled Exhibit and placeholders with letters such as Exhib xx when they are placed in square brackets. I refer to the SAMPLE INVOICE PDF dated 2 August 2026 [Exhib xx].",
        "Several exhibits can be detected separately in the same paragraph, including different file types. I refer to the SAMPLE PROJECT REPORT in DOCX format dated 3 August 2026, the SAMPLE CLAIMANT EMAIL in EML format dated 4 August 2026 and the SAMPLE COST WORKBOOK in XLSX format dated 5 August 2026 [AH-xx; AH-xx; AH-xx].",
        "For an XLSX exhibit, sheet selection means choosing which worksheet tabs and detected cell ranges Microsoft Excel will print into the bundle. Tick the sheets that belong in the exhibit; unticked sheets are omitted. The preview shows the planned A4 pages, saved formula results are used without recalculation, and the source workbook is not edited.",
        "This guided project also supplies a PDF cover template and a PDF index template. The index template contains a complete three-column table with placeholder headings for the item number, exhibit description and bundle pages. The tool places the final index entries inside those columns.",
        "A repeated reference can use the already selected source instead of creating a duplicate exhibit. I refer again to the SAMPLE AGREEMENT dated 1 August 2026 [Exhibit]. Every match still requires human confirmation, and the source statement remains unchanged.",
        "The SAMPLE UNREFERENCED CHECKLIST is deliberately not cited. Use it to practise the warning shown when a user deliberately adds an exhibit that is not referred to in the witness statement.",
        "At the final-order stage, add an index heading and place exhibits beneath it. The heading appears in the index and becomes a parent PDF bookmark. Manual moves apply immediately; an automatic sort is only a preview until you choose Use this order, and existing headings are not silently discarded.",
        "If a bundle is split into volumes, every volume contains the complete index. The requested maximum includes the cover, repeated index and exhibit pages, and the default printed numbering continues across volumes unless the user deliberately chooses a different scheme.",
        "A supplied template may already contain a case or matter number, party names or other matter details. Exhibit Builder preserves the supplied PDF appearance but does not assume those details are correct: review the exact rendered PDF and confirm the detected matter details, visible placeholders and any disagreement between templates separately.",
        "The current automatic citation matching and local OCR are tested for English. Original PDF artwork and text remain in the source language, but translated placeholder words, non-English OCR and fully multilingual generated index text are not claimed in this release and require careful manual review.",
    ]
    for text in paragraphs:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.keep_together = True
        apply_number(paragraph, num_id)
        set_run_font(paragraph.add_run(text), 11, "000000")

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(9)
    note.paragraph_format.space_after = Pt(0)
    set_run_font(note.add_run("Review rule: "), 10, CAUTION, True)
    set_run_font(note.add_run("Several placeholders in one paragraph are detected separately, but each proposed source, worksheet selection, template detail, grouping and final order remains subject to human review."), 10, CAUTION)

    document.core_properties.title = "Exhibit Builder Guided Sample Witness Statement"
    document.core_properties.subject = "Instructional fixture for offline exhibit-bundle demonstrations"
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    document.core_properties.keywords = "guided sample, exhibit placeholders, demonstration"
    document.save(OUT / "01_GUIDED_SAMPLE_Witness_Statement.docx")


PDF_SPECS = [
    ("01_SAMPLE_Agreement.pdf", "SAMPLE AGREEMENT", "Document date", "1 August 2026", "Agreement reference: DEMO-AGREEMENT-01", "This sample agreement demonstrates a clear title and date match."),
    ("02_SAMPLE_Invoice.pdf", "SAMPLE INVOICE", "Invoice date", "2 August 2026", "Invoice number: DEMO-INVOICE-02", "This sample invoice demonstrates a placeholder with letters."),
    ("06_SAMPLE_Unreferenced_Checklist.pdf", "SAMPLE UNREFERENCED CHECKLIST", "Checklist date", "6 August 2026", "Practice purpose: manual addition", "This document is deliberately not referred to in the guided statement. Use it to practise Add an exhibit."),
]


def create_pdf(file_name, title, date_label, date_value, reference, explanation):
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(str(OUT / file_name), pagesize=A4, rightMargin=22 * mm, leftMargin=22 * mm, topMargin=20 * mm, bottomMargin=20 * mm, title=title, author="Exhibit Builder", subject="Guided sample exhibit")
    banner = ParagraphStyle("banner", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=10, leading=13, textColor=colors.HexColor("#7A351F"), alignment=TA_CENTER)
    title_style = ParagraphStyle("title", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=24, leading=28, textColor=colors.HexColor("#0B2545"), alignment=TA_LEFT, spaceAfter=8)
    subtitle = ParagraphStyle("subtitle", parent=styles["Normal"], fontName="Helvetica", fontSize=11, leading=15, textColor=colors.HexColor("#5C6670"), spaceAfter=14)
    body = ParagraphStyle("body", parent=styles["BodyText"], fontName="Helvetica", fontSize=11, leading=16, textColor=colors.HexColor("#182433"), spaceAfter=10)
    small = ParagraphStyle("small", parent=body, fontSize=9, leading=12, textColor=colors.HexColor("#5C6670"), alignment=TA_CENTER)
    story = [
        Table([[Paragraph("SAMPLE DOCUMENT - FOR EXHIBIT BUILDER DEMONSTRATION ONLY", banner)]], colWidths=[166 * mm], style=TableStyle([("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#FFF0E8")), ("BOX", (0, 0), (-1, -1), 1, colors.HexColor("#CC8C75")), ("TOPPADDING", (0, 0), (-1, -1), 9), ("BOTTOMPADDING", (0, 0), (-1, -1), 9)])),
        Spacer(1, 16 * mm),
        Paragraph(title, title_style),
        Paragraph("An intentionally obvious training document for the guided sample.", subtitle),
        Table([[Paragraph("FIELD", banner), Paragraph("VALUE", banner)], [date_label, date_value], ["Sample reference", reference]], colWidths=[48 * mm, 118 * mm], style=TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")), ("GRID", (0, 0), (-1, -1), 0.6, colors.HexColor("#B8C4D0")), ("FONTNAME", (0, 1), (-1, -1), "Helvetica"), ("FONTSIZE", (0, 1), (-1, -1), 10), ("TEXTCOLOR", (0, 1), (-1, -1), colors.HexColor("#182433")), ("VALIGN", (0, 0), (-1, -1), "TOP"), ("LEFTPADDING", (0, 0), (-1, -1), 8), ("RIGHTPADDING", (0, 0), (-1, -1), 8), ("TOPPADDING", (0, 0), (-1, -1), 8), ("BOTTOMPADDING", (0, 0), (-1, -1), 8)])),
        Spacer(1, 12 * mm),
        Paragraph(explanation, body),
        Table([[Paragraph("What to notice", banner)], [Paragraph("The filename, title and date deliberately repeat the words used in the guided witness statement. This makes the sample easy to understand while still requiring the reviewer to confirm every match.", body)]], colWidths=[166 * mm], style=TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")), ("BACKGROUND", (0, 1), (-1, 1), colors.HexColor("#F4F6F9")), ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor("#B8C4D0")), ("LEFTPADDING", (0, 0), (-1, -1), 10), ("RIGHTPADDING", (0, 0), (-1, -1), 10), ("TOPPADDING", (0, 0), (-1, -1), 9), ("BOTTOMPADDING", (0, 0), (-1, -1), 9)])),
        Spacer(1, 33 * mm),
        Paragraph("This file contains no real case, party or personal information.", small),
    ]
    doc.build(story)


def create_project_report():
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    banner = document.add_table(rows=1, cols=1)
    set_table_geometry(banner, [9360])
    shade(banner.cell(0, 0), "FFF0E8")
    p = banner.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("SAMPLE DOCUMENT - FOR EXHIBIT BUILDER DEMONSTRATION ONLY"), 9, "7A351F", True)
    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(5)
    set_run_font(title.add_run("SAMPLE PROJECT REPORT"), 23, INK, True)
    subtitle = document.add_paragraph()
    set_run_font(subtitle.add_run("DOCX evidence example | 3 August 2026"), 11, MUTED, False, True)
    document.add_paragraph("This deliberately simple Word document demonstrates that DOCX files can be supplied as exhibit evidence. Exhibit Builder extracts readable text locally and prints a bundle-safe representation without changing this source file.")
    table = document.add_table(rows=4, cols=2)
    set_table_geometry(table, [3000, 6360])
    rows = [("Field", "Sample value"), ("Report reference", "DEMO-REPORT-03"), ("Status", "Ready for guided review"), ("What to notice", "The statement repeats this title, date and file type so the proposed match is obvious." )]
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            if row_index == 0:
                shade(cell, LIGHT_BLUE)
            p = cell.paragraphs[0]
            set_run_font(p.add_run(value), 10, INK, row_index == 0)
    document.core_properties.title = "Sample Project Report"
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    document.save(OUT / "03_SAMPLE_Project_Report.docx")


def create_sample_email():
    content = "\r\n".join([
        "From: Guided Claimant <claimant@example.invalid>",
        "To: Guided Respondent <respondent@example.invalid>",
        "Date: Tue, 04 Aug 2026 10:30:00 +0100",
        "Subject: SAMPLE CLAIMANT EMAIL - guided exhibit",
        "MIME-Version: 1.0",
        "Content-Type: text/plain; charset=UTF-8",
        "",
        "SAMPLE DOCUMENT - FOR EXHIBIT BUILDER DEMONSTRATION ONLY",
        "",
        "This is an intentionally obvious EML evidence example dated 4 August 2026.",
        "Reference: DEMO-EMAIL-04.",
        "The file demonstrates that a saved email can be selected as an exhibit and rendered locally with its email headers.",
        "",
    ])
    (OUT / "04_SAMPLE_Claimant_Email.eml").write_text(content, encoding="utf-8", newline="")


def create_cover_template():
    path = OUT / "00_GUIDED_SAMPLE_Cover_Template.pdf"
    pdf = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4
    pdf.setTitle("Guided Sample Cover Template")
    pdf.setFillColor(colors.HexColor("#17365D"))
    pdf.rect(0, height - 88, width, 88, fill=1, stroke=0)
    pdf.setFillColor(colors.white)
    pdf.setFont("Helvetica-Bold", 11)
    pdf.drawString(48, height - 52, "GUIDED SAMPLE - PDF COVER TEMPLATE")
    pdf.setFillColor(colors.HexColor("#17365D"))
    pdf.setFont("Helvetica-Bold", 25)
    pdf.drawString(48, height - 170, "[PROJECT OR CASE TITLE]")
    pdf.setFont("Helvetica", 15)
    pdf.drawString(48, height - 210, "[WITNESS EXHIBIT BUNDLE]")
    pdf.setStrokeColor(colors.HexColor("#9FB3C8"))
    pdf.line(48, height - 235, width - 48, height - 235)
    pdf.setFillColor(colors.HexColor("#4F5968"))
    pdf.setFont("Helvetica", 10)
    pdf.drawString(48, 92, "Tutorial note: the supplied PDF appearance is preserved and fitted to A4 without cropping.")
    pdf.drawString(48, 76, "Review and separately confirm visible matter details and placeholders before building.")
    pdf.showPage()
    pdf.save()


def create_index_template():
    path = OUT / "00_GUIDED_SAMPLE_Index_Template.pdf"
    pdf = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4
    pdf.setTitle("Guided Sample Full Table Index Template")
    left, number_right, description_right, right = 54, 95, 450, 550
    top, header_bottom, bottom = 746, 684, 132
    pdf.setFillColor(colors.HexColor("#17365D"))
    pdf.rect(left, top, right - left, 42, fill=1, stroke=0)
    pdf.setFillColor(colors.white)
    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawCentredString((left + right) / 2, top + 14, "GUIDED SAMPLE - INDEX OF EXHIBITS")
    pdf.setFillColor(colors.HexColor("#1F4E79"))
    pdf.rect(left, header_bottom, right - left, top - header_bottom, fill=1, stroke=0)
    pdf.setFillColor(colors.white)
    pdf.setFont("Helvetica-Bold", 5.5)
    pdf.drawCentredString((left + number_right) / 2, 712, "{{ITEM NO.}}")
    pdf.setFont("Helvetica-Bold", 7)
    pdf.drawString(106, 712, "{{EXHIBIT DESCRIPTION - INSERTED AUTOMATICALLY}}")
    pdf.drawCentredString((description_right + right) / 2, 712, "{{BUNDLE PAGES}}")
    pdf.setStrokeColor(colors.HexColor("#8EA9C1"))
    pdf.setLineWidth(0.8)
    pdf.rect(left, bottom, right - left, top - bottom, fill=0, stroke=1)
    pdf.line(number_right, bottom, number_right, top)
    pdf.line(description_right, bottom, description_right, top)
    pdf.line(left, header_bottom, right, header_bottom)
    pdf.setStrokeColor(colors.HexColor("#D7E1EA"))
    row_line = 638
    while row_line > bottom:
        pdf.line(left, row_line, right, row_line)
        row_line -= 46
    pdf.setFillColor(colors.HexColor("#5C6670"))
    pdf.setFont("Helvetica", 8)
    pdf.drawString(left, 102, "Placeholder headings define the three fixed areas used by Exhibit Builder.")
    pdf.drawString(left, 88, "Review this exact PDF and confirm its visible matter details and placeholders before building.")
    pdf.showPage()
    pdf.save()


if __name__ == "__main__":
    for obsolete_name in (
        "03_SAMPLE_Claimant_Confirmation_Email.pdf",
        "04_SAMPLE_Claimant_Delivery_Email.pdf",
        "05_SAMPLE_Claimant_Payment_Email.pdf",
    ):
        obsolete_path = OUT / obsolete_name
        if obsolete_path.exists():
            obsolete_path.unlink()
    create_statement()
    for spec in PDF_SPECS:
        create_pdf(*spec)
    create_project_report()
    create_sample_email()
    create_cover_template()
    create_index_template()
